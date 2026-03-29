"""DOCX generation - professional lead magnet styling.

Design patterns from the Guide Creation Playbook:
- Cover page: label + rule + large title + subtitle + author
- Callout boxes: single-cell tables with colored backgrounds
- Comparison tables: red/green two-column with icons
- Scenario tables: colored situation + white recommendation
- Framework steps: heading + explanation + bullets + ACTION callout
- Closing page: THE BOTTOM LINE + recap + author
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

# ---------------------------------------------------------------------------
# Color palette (from playbook reference)
# ---------------------------------------------------------------------------
CLR_NEAR_BLACK = RGBColor(0x0F, 0x17, 0x2A)  # #0F172A - titles, headings
CLR_BODY = RGBColor(0x1E, 0x29, 0x3B)  # #1E293B - body text (slate-800)
CLR_DIM = RGBColor(0x64, 0x74, 0x8B)  # #64748B - footer, captions
CLR_BLUE = RGBColor(0x25, 0x63, 0xEB)  # #2563EB - H2, accents
CLR_BLUE_DARK = RGBColor(0x1E, 0x3A, 0x8A)  # #1E3A8A - blue callout text
CLR_AMBER_DARK = RGBColor(0x78, 0x35, 0x0F)  # #78350F - amber callout text
CLR_RED_DARK = RGBColor(0x99, 0x1B, 0x1B)  # #991B1B - bad column text
CLR_GREEN_DARK = RGBColor(0x16, 0x65, 0x34)  # #166534 - good column text

# Background hex codes (for cell shading)
BG_BLUE_CALLOUT = "EFF6FF"
BG_AMBER_CALLOUT = "FFFBEB"
BG_NEUTRAL_CALLOUT = "F8FAFC"
BG_RED_LIGHT = "FEF2F2"
BG_GREEN_LIGHT = "F0FDF4"
BG_LIGHT_GRAY = "F1F5F9"

FONT_NAME = "Calibri"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background color on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    """Set cell margins in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def _set_cell_border(cell, color: str = "E2E8F0", sz: int = 4) -> None:
    """Set thin borders on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "start", "end"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:color"), color)
        border.set(qn("w:space"), "0")
        borders.append(border)
    tc_pr.append(borders)


def _remove_table_borders(table: Table) -> None:
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    tbl_pr.append(borders)


def _add_run(
    paragraph,
    text: str,
    *,
    bold=False,
    italic=False,
    size: int | None = None,
    color: RGBColor | None = None,
    font_name: str = FONT_NAME,
):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _add_styled_paragraph(
    doc,
    text: str,
    *,
    size=11,
    color=CLR_BODY,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=6,
):
    """Add a paragraph with consistent styling.

    Note: space_before/space_after are in POINTS (not DXA/twips).
    Playbook spec: body text = 6pt after, H1 = 10pt after, H2 = 8pt after.
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)  # 16pt line height for 11pt text
    _add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _add_horizontal_rule(doc, color: str = "2563EB", thickness: int = 8) -> None:
    """Add a colored horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_light_rule(doc, color: str = "CBD5E1") -> None:
    """Add a light gray horizontal rule."""
    _add_horizontal_rule(doc, color=color, thickness=4)


def _add_bullet(doc, text: str) -> None:
    """Add a bullet point paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-14)
    _add_run(p, "\u2022  ", size=11, color=CLR_BLUE)
    _add_run(p, text, size=11, color=CLR_BODY)


def _add_callout_box(doc, label: str, content: str, bg_hex: str, text_color: RGBColor) -> None:
    """Create a single-cell callout table with colored background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    cell = table.cell(0, 0)
    _set_cell_shading(cell, bg_hex)
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    # Subtle matching border
    border_color = bg_hex[:2] + "CC" + bg_hex[4:]  # slightly darker
    _set_cell_border(cell, color=border_color, sz=2)

    # Clear default paragraph and write content
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _add_run(p, f"{label.upper()}:  ", bold=True, size=10, color=text_color)
    _add_run(p, content, size=10, color=text_color)


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------


def _render_cover(doc, guide_data: dict) -> None:
    """Render the cover page."""
    # Label line
    _add_styled_paragraph(
        doc,
        "FREE GUIDE | 2026 EDITION",
        size=10,
        color=CLR_BLUE,
        bold=True,
        space_before=36,
        space_after=4,
    )

    # Blue divider
    _add_horizontal_rule(doc)

    # Main title
    title = guide_data.get("guide_title", "Weekly Guide")
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(12)
    _add_run(p_title, title, bold=True, size=28, color=CLR_NEAR_BLACK)

    # Subtitle
    subtitle = guide_data.get("subtitle", "")
    if subtitle:
        _add_styled_paragraph(
            doc,
            subtitle,
            size=13,
            color=CLR_DIM,
            space_before=0,
            space_after=10,
        )

    # Light divider
    _add_light_rule(doc)

    # Author block
    author_name = guide_data.get("author_name", "Ziv Raviv")
    author_url = guide_data.get("author_url", "zivraviv.com")
    p_author = doc.add_paragraph()
    p_author.paragraph_format.space_before = Pt(12)
    p_author.paragraph_format.space_after = Pt(0)
    _add_run(p_author, f"By {author_name}", bold=True, size=11, color=CLR_NEAR_BLACK)

    p_url = doc.add_paragraph()
    p_url.paragraph_format.space_before = Pt(2)
    p_url.paragraph_format.space_after = Pt(0)
    _add_run(p_url, author_url, size=10, color=CLR_BLUE)

    # Page break after cover
    doc.add_page_break()


def _render_narrative(doc, section: dict) -> None:
    """Render a narrative section (heading + paragraphs + auto-bullets)."""
    title = section.get("title", "")
    content = section.get("content", "")

    if title:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(24)
        p_h.paragraph_format.space_after = Pt(12)
        _add_run(p_h, title, bold=True, size=20, color=CLR_NEAR_BLACK)

    for para_text in content.split("\n\n"):
        text = para_text.strip()
        if not text:
            continue
        # Auto-detect bullet lists
        if text.startswith("- ") or text.startswith("* "):
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith(("- ", "* ")):
                    line = line[2:]
                if line:
                    _add_bullet(doc, line)
        else:
            _add_styled_paragraph(doc, text)


def _render_callout(doc, section: dict) -> None:
    """Render a callout box (colored single-cell table)."""
    label = section.get("label", "KEY INSIGHT")
    content = section.get("content", "")
    style = section.get("callout_style", "blue")

    if style == "amber":
        bg, text_color = BG_AMBER_CALLOUT, CLR_AMBER_DARK
    elif style == "neutral":
        bg, text_color = BG_NEUTRAL_CALLOUT, CLR_BODY
    else:
        bg, text_color = BG_BLUE_CALLOUT, CLR_BLUE_DARK

    _add_callout_box(doc, label, content, bg, text_color)


def _render_quick_win(doc, section: dict) -> None:
    """Render a Quick Win section - interactive worksheet with empty fill-in table."""
    title = section.get("title", "Your 15-Minute Quick Win")
    instruction = section.get("instruction", "")
    headers = section.get("table_headers", [])
    num_rows = section.get("table_rows", 5)
    what_you_learn = section.get("what_you_learn", "")

    # Section heading
    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(24)
    p_h.paragraph_format.space_after = Pt(12)
    _add_run(p_h, title, bold=True, size=20, color=CLR_NEAR_BLACK)

    # Instruction paragraph
    if instruction:
        _add_styled_paragraph(doc, instruction, space_after=12)

    # Worksheet table: header row + blank data rows
    if headers:
        num_cols = len(headers)
        table = doc.add_table(rows=num_rows + 1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row - light blue background
        for col_idx, header_text in enumerate(headers):
            cell = table.cell(0, col_idx)
            _set_cell_shading(cell, BG_BLUE_CALLOUT)
            _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            _set_cell_border(cell, color="93C5FD", sz=4)
            cell.paragraphs[0].clear()
            _add_run(
                cell.paragraphs[0], header_text, bold=True, size=10, color=CLR_BLUE_DARK
            )

        # Data rows - white background, thin borders
        for row_idx in range(1, num_rows + 1):
            for col_idx in range(num_cols):
                cell = table.cell(row_idx, col_idx)
                _set_cell_shading(cell, "FFFFFF")
                _set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
                _set_cell_border(cell, color="CBD5E1", sz=2)
                cell.paragraphs[0].clear()
                # Empty cell - reader fills it in
                _add_run(cell.paragraphs[0], " ", size=10, color=CLR_BODY)

    # "What you'll learn" callout
    if what_you_learn:
        _add_callout_box(
            doc, "WHAT YOU'LL LEARN", what_you_learn, BG_AMBER_CALLOUT, CLR_AMBER_DARK
        )


def _render_comparison(doc, section: dict) -> None:
    """Render a red/green comparison table."""
    title = section.get("title", "")
    if title:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(24)
        p_h.paragraph_format.space_after = Pt(12)
        _add_run(p_h, title, bold=True, size=20, color=CLR_NEAR_BLACK)

    bad_label = section.get("bad_label", "Before")
    bad_items = section.get("bad_items", [])
    good_label = section.get("good_label", "After")
    good_items = section.get("good_items", [])

    max_rows = max(len(bad_items), len(good_items), 1)
    table = doc.add_table(rows=max_rows + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, (label, bg, text_color, icon) in enumerate(
        [
            (bad_label, BG_RED_LIGHT, CLR_RED_DARK, "\u2717"),
            (good_label, BG_GREEN_LIGHT, CLR_GREEN_DARK, "\u2713"),
        ]
    ):
        cell = table.cell(0, i)
        _set_cell_shading(cell, bg)
        _set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        _set_cell_border(cell, color="E2E8F0", sz=2)
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], f"{icon} {label}", bold=True, size=12, color=text_color)

    # Item rows
    items_lists = [bad_items, good_items]
    colors = [CLR_RED_DARK, CLR_GREEN_DARK]
    bgs = [BG_RED_LIGHT, BG_GREEN_LIGHT]
    icons = ["\u2717", "\u2713"]

    for row_idx in range(max_rows):
        for col_idx in range(2):
            cell = table.cell(row_idx + 1, col_idx)
            _set_cell_shading(cell, bgs[col_idx])
            _set_cell_margins(cell, top=40, bottom=40, left=120, right=120)
            _set_cell_border(cell, color="E2E8F0", sz=2)
            cell.paragraphs[0].clear()
            items = items_lists[col_idx]
            if row_idx < len(items):
                _add_run(
                    cell.paragraphs[0],
                    f"{icons[col_idx]} {items[row_idx]}",
                    size=10,
                    color=colors[col_idx],
                )


def _render_framework(doc, section: dict) -> None:
    """Render a framework section with numbered steps and ACTION callouts."""
    title = section.get("title", "")
    if title:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(24)
        p_h.paragraph_format.space_after = Pt(12)
        _add_run(p_h, title, bold=True, size=20, color=CLR_NEAR_BLACK)

    intro = section.get("intro", "")
    if intro:
        _add_styled_paragraph(doc, intro)

    for i, step in enumerate(section.get("steps", []), 1):
        label = step.get("label", f"Step {i}")
        explanation = step.get("explanation", "")
        action = step.get("action", "")
        bullets = step.get("bullets", [])

        # Step heading: blue number + dark label
        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_before = Pt(18)
        p_step.paragraph_format.space_after = Pt(8)
        _add_run(p_step, f"{i}. ", bold=True, size=15, color=CLR_BLUE)
        _add_run(p_step, label, bold=True, size=15, color=CLR_NEAR_BLACK)

        if explanation:
            _add_styled_paragraph(doc, explanation)

        for bullet in bullets:
            _add_bullet(doc, bullet)

        if action:
            _add_callout_box(doc, "ACTION", action, BG_BLUE_CALLOUT, CLR_BLUE_DARK)


def _render_scenarios(doc, section: dict) -> None:
    """Render scenario tables - colored situation + white recommendation."""
    title = section.get("title", "")
    if title:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(24)
        p_h.paragraph_format.space_after = Pt(12)
        _add_run(p_h, title, bold=True, size=20, color=CLR_NEAR_BLACK)

    intro = section.get("intro", "")
    if intro:
        _add_styled_paragraph(doc, intro)

    accent_bgs = [BG_BLUE_CALLOUT, BG_AMBER_CALLOUT, BG_LIGHT_GRAY]
    accent_colors = [CLR_BLUE_DARK, CLR_AMBER_DARK, CLR_BODY]

    for idx, scenario in enumerate(section.get("scenarios", [])):
        situation = scenario.get("situation", "")
        response = scenario.get("response", "")

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        bg_idx = idx % len(accent_bgs)

        # Left - situation (colored)
        left = table.cell(0, 0)
        _set_cell_shading(left, accent_bgs[bg_idx])
        _set_cell_margins(left, top=80, bottom=80, left=120, right=120)
        _set_cell_border(left, color="E2E8F0", sz=2)
        left.paragraphs[0].clear()
        _add_run(
            left.paragraphs[0],
            f'"{situation}"',
            bold=True,
            italic=True,
            size=10,
            color=accent_colors[bg_idx],
        )

        # Right - recommendation (white)
        right = table.cell(0, 1)
        _set_cell_shading(right, "FFFFFF")
        _set_cell_margins(right, top=80, bottom=80, left=120, right=120)
        _set_cell_border(right, color="E2E8F0", sz=2)
        right.paragraphs[0].clear()
        _add_run(right.paragraphs[0], response, size=10, color=CLR_BODY)


def _render_closing(doc, section: dict, guide_data: dict) -> None:
    """Render the closing page - mirrors cover structure."""
    doc.add_page_break()

    _add_styled_paragraph(
        doc,
        "THE BOTTOM LINE",
        size=10,
        color=CLR_BLUE,
        bold=True,
        space_before=60,
        space_after=8,
    )
    _add_horizontal_rule(doc)

    headline = section.get("headline", "")
    if headline:
        # Short headlines (under 80 chars) get large treatment; longer ones get body size
        is_short = len(headline) < 80
        p_big = doc.add_paragraph()
        p_big.paragraph_format.space_before = Pt(20)
        p_big.paragraph_format.space_after = Pt(16)
        _add_run(p_big, headline, bold=True, size=18 if is_short else 13, color=CLR_NEAR_BLACK)

    # "What You Now Have" - concrete outputs the reader produced
    you_now_have = section.get("you_now_have", [])
    if you_now_have:
        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(12)
        p_label.paragraph_format.space_after = Pt(6)
        _add_run(p_label, "WHAT YOU NOW HAVE:", bold=True, size=11, color=CLR_BLUE)

        for item_text in you_now_have:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-14)
            _add_run(p, "\u2713  ", size=11, color=CLR_BLUE)
            _add_run(p, item_text, size=11, color=CLR_NEAR_BLACK)

    # Fallback: legacy recap_steps (backward compat)
    recap_steps = section.get("recap_steps", [])
    if recap_steps and not you_now_have:
        for i, step_text in enumerate(recap_steps, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_run(p, f"{i}. ", bold=True, size=12, color=CLR_BLUE)
            _add_run(p, step_text, size=12, color=CLR_NEAR_BLACK)

    _add_light_rule(doc)

    author_name = guide_data.get("author_name", "Ziv Raviv")
    author_url = guide_data.get("author_url", "zivraviv.com")

    p_author = doc.add_paragraph()
    p_author.paragraph_format.space_before = Pt(16)
    p_author.paragraph_format.space_after = Pt(4)
    _add_run(p_author, author_name, bold=True, size=11, color=CLR_NEAR_BLACK)

    p_url = doc.add_paragraph()
    p_url.paragraph_format.space_before = Pt(0)
    p_url.paragraph_format.space_after = Pt(8)
    _add_run(p_url, author_url, size=10, color=CLR_BLUE)

    cta = section.get("cta", "")
    if cta:
        _add_styled_paragraph(doc, cta, size=11, color=CLR_DIM, italic=True, space_before=4)


def _setup_footer(doc, author_name: str, author_url: str) -> None:
    """Add footer with page number and author info."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    _add_run(p, f"By {author_name}  |  {author_url}  |  Page ", size=8, color=CLR_DIM)

    # Page number field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1 = p.add_run()
    r1._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2 = p.add_run()
    r2.font.size = Pt(8)
    r2.font.color.rgb = CLR_DIM
    r2._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3 = p.add_run()
    r3._r.append(fld_end)


# ---------------------------------------------------------------------------
# Section dispatcher
# ---------------------------------------------------------------------------

SECTION_RENDERERS = {
    "narrative": _render_narrative,
    "callout": _render_callout,
    "quick_win": _render_quick_win,
    "comparison": _render_comparison,
    "framework": _render_framework,
    "scenarios": _render_scenarios,
}


# ---------------------------------------------------------------------------
# Post-processing helpers
# ---------------------------------------------------------------------------


def _clamp_paragraph_spacing(doc) -> None:
    """Cap all paragraph space_after to 24pt (480 twips) max.

    python-docx can sometimes produce oversized spacing values due to
    unit confusion or style inheritance.  This safety net ensures no
    body paragraph ends up with e.g. 120pt of trailing whitespace.
    """
    max_twips = 480  # 24pt

    for para in doc.element.body.iterchildren(qn("w:p")):
        p_pr = para.find(qn("w:pPr"))
        if p_pr is None:
            continue
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            continue
        after_val = spacing.get(qn("w:after"))
        if after_val and after_val.isdigit() and int(after_val) > max_twips:
            spacing.set(qn("w:after"), str(max_twips))
        before_val = spacing.get(qn("w:before"))
        if before_val and before_val.isdigit() and int(before_val) > 960:  # 48pt max
            spacing.set(qn("w:before"), "960")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def create_guide_docx(
    guide_data: dict[str, Any] | str,
    sections_or_path: list[dict[str, str]] | str = "",
    output_path: str = "",
    author_name: str = "Ziv Raviv",
    author_url: str = "zivraviv.com",
) -> str:
    """Create a professionally styled DOCX guide.

    Supports two calling conventions for backward compatibility:
        - New: create_guide_docx(guide_data_dict, output_path_str)
        - Legacy: create_guide_docx(title_str, sections_list, output_path_str)

    Returns:
        The output path.
    """
    # Detect legacy calling convention: create_guide_docx("Title", [...], "path")
    if isinstance(guide_data, str):
        title = guide_data
        sections = sections_or_path if isinstance(sections_or_path, list) else []
        out_path = (
            output_path
            if output_path
            else (sections_or_path if isinstance(sections_or_path, str) else "")
        )
        guide_data = {
            "guide_title": title,
            "subtitle": "",
            "author_name": author_name,
            "author_url": author_url,
            "sections": [
                {"type": "narrative", "title": s.get("title", ""), "content": s.get("content", "")}
                for s in sections
            ],
        }
    else:
        out_path = sections_or_path if isinstance(sections_or_path, str) else output_path
        guide_data.setdefault("author_name", author_name)
        guide_data.setdefault("author_url", author_url)

    doc = docx.Document()

    # Global styles
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = CLR_BODY
    style.paragraph_format.space_after = Pt(6)

    # Page margins
    sec_fmt = doc.sections[0]
    sec_fmt.top_margin = Inches(1)
    sec_fmt.bottom_margin = Inches(1)
    sec_fmt.left_margin = Inches(1)
    sec_fmt.right_margin = Inches(1)

    # Footer
    _setup_footer(
        doc,
        guide_data.get("author_name", "Ziv Raviv"),
        guide_data.get("author_url", "zivraviv.com"),
    )

    # Cover page
    _render_cover(doc, guide_data)

    # Sections
    for sec in guide_data.get("sections", []):
        sec_type = sec.get("type", "narrative")
        if sec_type == "closing":
            _render_closing(doc, sec, guide_data)
        elif sec_type in SECTION_RENDERERS:
            SECTION_RENDERERS[sec_type](doc, sec)
        else:
            _render_narrative(doc, sec)

    # Clamp spacing - prevent runaway space_after values (>24pt = 480 twips)
    _clamp_paragraph_spacing(doc)

    # Save
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
