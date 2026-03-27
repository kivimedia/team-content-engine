"""Document ingestion service - parses uploaded files into SourceDocument records."""

from __future__ import annotations

import asyncio
import base64
from typing import Any

from datetime import datetime, timezone

import anthropic
import docx
import structlog

from sqlalchemy.ext.asyncio import AsyncSession

from tce.models.source_document import SourceDocument
from tce.settings import settings

logger = structlog.get_logger()


class DocumentIngestService:
    """Handle document uploads and text extraction."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def ingest_docx(self, file_path: str, file_name: str) -> dict[str, Any]:
        """Parse a DOCX file and create a SourceDocument record."""
        doc = docx.Document(file_path)

        # Extract all text
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # GAP-02: Extract and OCR embedded images
        image_texts = await self._extract_images_ocr(doc)
        if image_texts:
            full_text += "\n\n## OCR from Embedded Images\n\n" + "\n\n".join(image_texts)

        # Count pages (approximate from paragraph count)
        approx_pages = max(1, len(paragraphs) // 20)

        # GAP-04: Upload source DOCX to S3 if configured
        s3_path = None
        try:
            from tce.services.storage import StorageService

            storage = StorageService()
            if storage.configured:
                import uuid as _uuid

                key = f"corpus/{_uuid.uuid4().hex}/{file_name}"
                with open(file_path, "rb") as f:
                    s3_path = await storage.upload(f.read(), key, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                logger.info("ingest.s3_uploaded", key=key)
        except Exception:
            logger.exception("ingest.s3_upload_failed")

        # Create source document record
        source_doc = SourceDocument(
            file_name=file_name,
            file_type="docx",
            language="he",  # Hebrew corpus default
            pages=approx_pages,
            extracted_text=full_text,
            ingested_at=datetime.now(timezone.utc),
            metadata_={
                "paragraph_count": len(paragraphs),
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "images_ocr_count": len(image_texts),
                "s3_path": s3_path,
            },
        )
        self.db.add(source_doc)
        await self.db.flush()

        return {
            "document_id": str(source_doc.id),
            "document_text": full_text,
            "file_name": file_name,
            "paragraph_count": len(paragraphs),
            "approx_pages": approx_pages,
            "images_ocr_count": len(image_texts),
        }

    async def _extract_images_ocr(self, doc: docx.Document) -> list[str]:
        """Extract embedded images from DOCX and OCR them via Claude vision."""
        if not settings.anthropic_api_key.get_secret_value():
            return []

        image_texts = []
        try:
            client = anthropic.AsyncAnthropic(
                api_key=settings.anthropic_api_key.get_secret_value()
            )
            for rel in doc.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                image_data = rel.target_part.blob
                content_type = rel.target_part.content_type or "image/png"

                # Skip very small images (likely icons/bullets)
                if len(image_data) < 5000:
                    continue

                # Per-image try/catch so one failure doesn't skip the rest
                try:
                    b64 = base64.standard_b64encode(image_data).decode("utf-8")
                    response = await asyncio.wait_for(
                        client.messages.create(
                            model=settings.haiku_model,
                            max_tokens=2000,
                            messages=[{
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image",
                                        "source": {
                                            "type": "base64",
                                            "media_type": content_type,
                                            "data": b64,
                                        },
                                    },
                                    {
                                        "type": "text",
                                        "text": (
                                            "Extract ALL text from this image. "
                                            "This may be a screenshot of a social media post. "
                                            "Include: post text, comments count, shares count, "
                                            "engagement metrics, author name. "
                                            "If text is in Hebrew, transcribe it in Hebrew. "
                                            "Return only the extracted text, no commentary."
                                        ),
                                    },
                                ],
                            }],
                        ),
                        timeout=30,
                    )

                    for block in response.content:
                        if block.type == "text" and block.text.strip():
                            image_texts.append(block.text.strip())

                    logger.info("ocr.image_processed", size=len(image_data))
                except asyncio.TimeoutError:
                    logger.warning("ocr.image_timeout", size=len(image_data))
                except Exception:
                    logger.exception("ocr.single_image_failed", size=len(image_data))

        except Exception:
            logger.exception("ocr.extraction_failed")

        return image_texts

    async def ingest_text(self, text: str, file_name: str) -> dict[str, Any]:
        """Ingest raw text content."""
        source_doc = SourceDocument(
            file_name=file_name,
            file_type="text",
            language="en",
            extracted_text=text,
            ingested_at=datetime.now(timezone.utc),
            metadata_={
                "char_count": len(text),
                "word_count": len(text.split()),
            },
        )
        self.db.add(source_doc)
        await self.db.flush()

        return {
            "document_id": str(source_doc.id),
            "document_text": text,
            "file_name": file_name,
        }
